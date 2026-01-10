from __future__ import annotations

import sys
import tempfile
from pathlib import Path

import httpx
from docx import Document


def main() -> int:
    api_base = sys.argv[1] if len(sys.argv) > 1 else "http://127.0.0.1:8000"
    url = api_base.rstrip("/") + "/api/convert"

    with tempfile.TemporaryDirectory() as td:
        docx_path = Path(td) / "smoke.docx"
        doc = Document()
        doc.add_heading("Smoke Test", level=1)
        doc.add_paragraph("Hello from smoke test.")
        doc.save(str(docx_path))

        with docx_path.open("rb") as f:
            files = {
                "file": (
                    docx_path.name,
                    f,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            }
            res = httpx.post(url, files=files, timeout=30)

    print("status=", res.status_code)
    try:
        body = res.json()
    except Exception:
        print(res.text)
        return 2

    print(body)
    if res.status_code != 200:
        return 1
    if not body.get("ok"):
        return 1
    data = body.get("data") or {}
    output = data.get("output") or {}
    if not (output.get("html") or output.get("text")):
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
